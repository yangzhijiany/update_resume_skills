# -*- coding: utf-8 -*-
"""
Update SKILLS section in Word resume (v8, user configurable style)
- Works for paragraphs in body and inside tables
- Removes old content after SKILLS until blank line or next heading (keeps heading itself)
- Inserts three skill lines with configurable font, size, spacing
Requires: pip install python-docx
"""

import os
from typing import Dict, List, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from dotenv import load_dotenv
from docx2pdf import convert

# ---------- Configurable style ----------
load_dotenv()

# font name can stay as string
FONT_NAME = os.getenv("FONT_NAME", "Times New Roman")

# convert to float
FONT_SIZE_PT = float(os.getenv("FONT_SIZE_PT", 10.5))

# convert to float (pt)
SPACE_AFTER_PT = float(os.getenv("SPACE_AFTER_PT", 6))

# handle line spacing rule (string → enum)
line_rule = os.getenv("LINE_SPACING_RULE", "SINGLE").upper()
if line_rule == "SINGLE":
    LINE_SPACING_RULE = WD_LINE_SPACING.SINGLE
elif line_rule == "DOUBLE":
    LINE_SPACING_RULE = WD_LINE_SPACING.DOUBLE
elif line_rule == "ONE_POINT_FIVE":
    LINE_SPACING_RULE = WD_LINE_SPACING.ONE_POINT_FIVE
else:
    LINE_SPACING_RULE = WD_LINE_SPACING.SINGLE  # fallback


# ---------- Utilities ----------

def iter_all_paragraphs(doc: Document):
    """Iterate all paragraphs in the document, including those inside tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def looks_like_heading(p: Paragraph) -> bool:
    """Check if a paragraph looks like a section heading."""
    txt = (p.text or "").strip()
    style_name = getattr(getattr(p, "style", None), "name", "") or ""
    if "Heading" in style_name:
        return True
    if txt and txt.isupper() and len(txt) <= 40:
        return True
    if txt.endswith(":") and txt[:-1].isupper() and len(txt) <= 41:
        return True
    return False


def iter_following_paragraphs(skills_para: Paragraph):
    """Yield following sibling paragraphs after the SKILLS heading."""
    el = skills_para._p.getnext()
    while el is not None:
        if el.tag.endswith("}p"):
            yield Paragraph(el, skills_para._parent)
        el = el.getnext()


def get_content_style_after(skills_para: Paragraph):
    """Get style of the first non-empty paragraph after SKILLS (to reuse for formatting)."""
    for p in iter_following_paragraphs(skills_para):
        txt = (p.text or "").strip()
        if not txt:
            continue
        if looks_like_heading(p):
            break
        return p.style
    return None


def force_run_font(run):
    """Force a run to use configured font and size."""
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)

    r = run._element
    rPr = r.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.append(rPr)
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def insert_paragraph_after(paragraph: Paragraph, text: str = "",
                           bold_label: bool = False, style=None) -> Paragraph:
    """Insert a new paragraph after the given one, with formatting applied."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass

    pf = new_para.paragraph_format
    pf.space_after = Pt(SPACE_AFTER_PT)
    pf.line_spacing_rule = LINE_SPACING_RULE

    if text:
        if bold_label and ":" in text:
            label, rest = text.split(":", 1)
            r1 = new_para.add_run(label + ":")
            r1.bold = True
            force_run_font(r1)

            r2 = new_para.add_run(rest)
            force_run_font(r2)
        elif bold_label:
            r = new_para.add_run(text)
            r.bold = True
            force_run_font(r)
        else:
            r = new_para.add_run(text)
            force_run_font(r)
    return new_para


def delete_block_after(skills_para: Paragraph):
    """Delete content after SKILLS until a blank line (remove it too) or the next heading."""
    el = skills_para._p.getnext()
    while el is not None:
        if not el.tag.endswith("}p"):
            el = el.getnext()
            continue

        p = Paragraph(el, skills_para._parent)
        txt = (p.text or "").strip()

        if not txt:
            nxt = el.getnext()
            p._element.getparent().remove(el)
            break

        if looks_like_heading(p):
            break

        nxt = el.getnext()
        p._element.getparent().remove(el)
        el = nxt


# ---------- Main ----------

def update_resume_skills(input_path: str, output_path: str, skills: Dict[str, List[str]]) -> None:
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {input_path}")

    doc = Document(input_path)

    skills_para: Optional[Paragraph] = None
    for p in iter_all_paragraphs(doc):
        if (p.text or "").strip().upper() == "SKILLS":
            skills_para = p
            break
    if not skills_para:
        raise ValueError("Could not find SKILLS heading (must be exactly 'SKILLS')")

    content_style = get_content_style_after(skills_para)
    if content_style is None:
        try:
            content_style = doc.styles["Normal"]
        except Exception:
            content_style = None

    delete_block_after(skills_para)

    after = insert_paragraph_after(
        skills_para,
        f"Programming & Frameworks: {', '.join(skills.get('programming', []))}",
        bold_label=True,
        style=content_style,
    )
    after = insert_paragraph_after(
        after,
        f"Software Development: {', '.join(skills.get('development', []))}",
        bold_label=True,
        style=content_style,
    )
    insert_paragraph_after(
        after,
        f"AI & Data Science: {', '.join(skills.get('ai', []))}",
        bold_label=True,
        style=content_style,
    )

    # Save DOCX
    doc.save(output_path)
    print(f"Resume updated, saved as: {output_path}")

    # Export to PDF
    pdf_path = output_path.replace(".docx", ".pdf")
    try:
        convert(output_path, pdf_path)
        print(f"PDF exported: {pdf_path}")
    except Exception as e:
        print(f"PDF export failed: {e}")


# ---------- Example run ----------
if __name__ == "__main__":
    input_file = r"E:\fetch_skills\Zhijian Yang - Resume.docx"
    output_file = r"E:\fetch_skills\resume_updated.docx"

    skills_dict = {
        "programming": ["Java", "C/C++", "Python", "Vue.js", "React", "Docker", "Git"],
        "development": ["MySQL", "MongoDB", "Firebase", "Spring Boot", "Redis"],
        "ai": ["RAG", "R", "Pandas", "Regression analysis", "A/B Testing", "Tableau"]
    }

    update_resume_skills(input_file, output_file, skills_dict)
